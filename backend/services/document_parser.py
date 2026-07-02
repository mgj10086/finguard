"""
文档解析服务

支持 PDF 和 Word 文档解析，提取文本、表格、章节结构。
利用 Celery 异步执行耗时解析任务。
"""

import logging
import os
import re
from pathlib import Path
from typing import Optional

from backend.core.celery_app import celery_app

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器 — PDF + Word"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    def parse(self, file_path: str) -> dict:
        """
        解析文档，返回结构化内容

        Returns:
            {
                "text": "纯文本内容",
                "pages": int,
                "sections": [{"title": "...", "content": "..."}],
                "tables": [{"caption": "...", "data": [[...]]}],
                "metadata": {...}
            }
        """
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext == ".docx":
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _parse_pdf(self, file_path: str) -> dict:
        """解析 PDF 文件"""
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        sections = []
        full_text = []
        tables = []

        for page_num, page in enumerate(doc):
            text = page.get_text()
            full_text.append(text)

            # 简单章节检测：匹配"第X章"、"第X节"等标题
            lines = text.split("\n")
            for line in lines:
                line = line.strip()
                if re.match(r"^第[一二三四五六七八九十百千]+[章节条]", line):
                    sections.append({
                        "title": line,
                        "content": "",
                        "page": page_num + 1,
                    })

            # 表格检测（PyMuPDF 提取）
            page_tables = page.find_tables()
            for tab in page_tables:
                table_data = tab.extract()
                if table_data:
                    tables.append({
                        "page": page_num + 1,
                        "data": table_data,
                    })

        # 给章节填充内容
        for i, section in enumerate(sections):
            start = full_text[section["page"] - 1].find(section["title"])
            if i < len(sections) - 1:
                end = full_text[section["page"] - 1].find(sections[i + 1]["title"])
                section["content"] = full_text[section["page"] - 1][start:end] if end > start else ""
            else:
                section["content"] = full_text[section["page"] - 1][start:]

        result = {
            "text": "\n".join(full_text),
            "pages": len(doc),
            "sections": sections,
            "tables": tables,
            "metadata": {
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "page_count": len(doc),
            },
        }
        doc.close()
        return result

    def _parse_docx(self, file_path: str) -> dict:
        """解析 Word 文件"""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        sections = []
        full_text = []
        tables = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            full_text.append(text)

            # 章节标题检测（根据样式或正则）
            style_name = para.style.name.lower() if para.style else ""
            is_heading = "heading" in style_name or re.match(
                r"^第[一二三四五六七八九十百千]+[章节条]", text
            )
            if is_heading:
                current_section = {"title": text, "content": ""}
                sections.append(current_section)
            elif current_section is not None:
                current_section["content"] += text + "\n"

        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append({"data": table_data})

        return {
            "text": "\n".join(full_text),
            "pages": 0,  # Word 文档无页数概念
            "sections": sections,
            "tables": tables,
            "metadata": {
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "paragraphs": len(doc.paragraphs),
            },
        }


# ===== Celery 异步任务 =====

@celery_app.task(bind=True, name="parse_document")
def parse_document_task(self, file_path: str) -> dict:
    """
    Celery 异步文档解析任务

    Args:
        file_path: 文档绝对路径

    Returns:
        解析结果 dict
    """
    self.update_state(state="PARSING")
    parser = DocumentParser()
    try:
        result = parser.parse(file_path)
        self.update_state(state="PARSED")
        return result
    except Exception as e:
        logger.error(f"Document parsing failed: {e}", exc_info=True)
        self.update_state(state="FAILED")
        raise
